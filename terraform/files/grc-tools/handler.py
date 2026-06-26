"""
grc-tools Lambda handler — policy builder backend.

Route dispatch for all /api/grc-tools/* endpoints.
Deployed as a standard zip-based Lambda (pure Python deps only).

"""

import json
import os
import io
import zipfile
import traceback
from datetime import datetime, timezone

import boto3
import jwt
import markdown
from jinja2 import Environment, FileSystemLoader, TemplateNotFound

# PDF generation (xhtml2pdf -- pure Python, no system deps)
from xhtml2pdf import pisa
from io import BytesIO

# DOCX generation
from docx import Document
from docx.shared import Pt, Inches  # noqa: F401

# --- Configuration ----------------------------------------------------------

SESSIONS_TABLE = os.environ["SESSIONS_TABLE"]
POLICIES_TABLE = os.environ["POLICIES_TABLE"]
S3_BUCKET = os.environ["S3_BUCKET"]
S3_PREFIX = os.environ["S3_PREFIX"]
USER_POOL_ID = os.environ["USER_POOL_ID"]
CLIENT_ID = os.environ["CLIENT_ID"]
ORIGIN_SECRET = os.environ["ORIGIN_SECRET"]
REGION = "us-east-1"

dynamodb = boto3.client("dynamodb")
s3 = boto3.client("s3")

# Jinja2 environment pointing at bundled templates
TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "policy-templates")
jinja_env = Environment(loader=FileSystemLoader(TEMPLATE_DIR))

# Load manifest
_manifest_path = os.path.join(os.path.dirname(__file__), "manifest.json")
with open(_manifest_path, "r") as f:
    MANIFEST = json.load(f)


# --- Helpers ----------------------------------------------------------------

def parse_cookies(headers):
    """Extract cookies from API Gateway v2 event headers."""
    cookies = {}
    cookie_header = headers.get("cookie", "")
    if cookie_header:
        for part in cookie_header.split(";"):
            part = part.strip()
            if "=" in part:
                k, v = part.split("=", 1)
                cookies[k.strip()] = v.strip()
    return cookies


def get_user_id(event):
    """Extract Cognito user_id from the grc_session JWT cookie."""
    cookies = parse_cookies(event.get("headers", {}))
    token = cookies.get("grc_session")
    if not token:
        return None

    # Decode without verification — the Lambda@Edge already verified it.
    # We just need the sub claim. Still check expiry.
    try:
        payload = jwt.decode(token, options={"verify_signature": False})
        if payload.get("exp", 0) < datetime.now(timezone.utc).timestamp():
            return None
        # Audience check
        aud = payload.get("aud") or payload.get("client_id")
        if aud != CLIENT_ID:
            return None
        return payload.get("sub")
    except Exception:
        return None


def check_origin_secret(event):
    """Validate the X-Origin-Secret header injected by CloudFront."""
    headers = event.get("headers", {})
    secret = headers.get("x-origin-secret", "")
    return secret == ORIGIN_SECRET


def json_response(body, status=200):
    return {
        "statusCode": status,
        "headers": {
            "Content-Type": "application/json",
            "Access-Control-Allow-Origin": "https://brooks-security.com",
            "Access-Control-Allow-Credentials": "true",
        },
        "body": json.dumps(body),
    }


def error_response(message, status=400):
    return json_response({"error": message}, status)


def get_session(user_id):
    """Read wizard state from DynamoDB."""
    resp = dynamodb.get_item(
        TableName=SESSIONS_TABLE,
        Key={"user_id": {"S": user_id}},
    )
    item = resp.get("Item")
    if not item:
        return None
    return deserialize_session(item)


def put_session(user_id, state):
    """Write wizard state to DynamoDB."""
    now = datetime.now(timezone.utc).isoformat()
    item = serialize_session(user_id, state, now)
    dynamodb.put_item(TableName=SESSIONS_TABLE, Item=item)
    return now


def deserialize_session(item):
    """Convert DynamoDB item to plain dict."""
    state = {}
    for k, v in item.items():
        if "S" in v:
            state[k] = v["S"]
        elif "N" in v:
            state[k] = int(v["N"])
        elif "SS" in v:
            state[k] = list(v["SS"])
        elif "NS" in v:
            state[k] = [int(n) for n in v["NS"]]
    return state


def serialize_session(user_id, state, updated_at):
    """Convert plain dict to DynamoDB item."""
    item = {
        "user_id": {"S": user_id},
        "updated_at": {"S": updated_at},
    }
    for k, v in state.items():
        if v is None:
            continue
        if isinstance(v, list) and all(isinstance(x, str) for x in v):
            item[k] = {"SS": v}
        elif isinstance(v, list) and all(isinstance(x, int) for x in v):
            item[k] = {"NS": [str(x) for x in v]}
        elif isinstance(v, int):
            item[k] = {"N": str(v)}
        elif isinstance(v, str):
            item[k] = {"S": v}
    return item


def extract_user_info(token):
    """Extract user info from JWT without verification (already verified by Lambda@Edge)."""
    try:
        payload = jwt.decode(token, options={"verify_signature": False})
        return {
            "user_id": payload.get("sub"),
            "email": payload.get("email", ""),
            "name": payload.get("name", ""),
            "idp": payload.get("identities", "[]"),
        }
    except Exception:
        return None


def parse_idp(identities_json_str):
    """Parse the Cognito identities JSON to extract provider name."""
    try:
        identities = json.loads(identities_json_str)
        if identities and len(identities) > 0:
            provider = identities[0].get("providerName", "")
            if "Google" in provider:
                return "Google"
            elif "Microsoft" in provider:
                return "Microsoft"
        return "Unknown"
    except (json.JSONDecodeError, IndexError):
        return "Unknown"


# --- PDF / DOCX Generators -------------------------------------------------

PDF_CSS = """
@page { size: letter; margin: 1in; }
body { font-family: 'DejaVu Sans', sans-serif; font-size: 11pt; line-height: 1.6; color: #24292f; }
h1 { font-size: 18pt; color: #0a66c2; border-bottom: 2px solid #0a66c2; padding-bottom: 6pt; margin-top: 0; }
h2 { font-size: 14pt; color: #24292f; margin-top: 18pt; }
h3 { font-size: 12pt; }
table { border-collapse: collapse; width: 100%; margin: 12pt 0; }
th, td { border: 1px solid #d0d7de; padding: 6pt 8pt; text-align: left; }
th { background: #f6f8fa; font-weight: 600; }
code { background: #f6f8fa; padding: 1pt 4pt; border-radius: 3pt; font-size: 10pt; }
pre { background: #f6f8fa; padding: 8pt; border-radius: 4pt; overflow-x: auto; }
strong { color: #24292f; }
"""


def render_pdf(md_content: str) -> bytes:
    """Render markdown to PDF via xhtml2pdf (pure Python)."""
    html = markdown.markdown(md_content, extensions=["tables", "fenced_code"])
    full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>{PDF_CSS}</style></head>
<body>{html}</body></html>"""
    buf = BytesIO()
    pisa.CreatePDF(full_html, dest=buf)
    return buf.getvalue()


def render_docx(md_content: str) -> bytes:
    """Render markdown to DOCX via python-docx."""
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = md_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)

        # Tables
        elif line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            i -= 1  # back up, will increment

            if len(table_lines) >= 2:
                # Skip separator row (index 1)
                headers = [c.strip() for c in table_lines[0].split("|")[1:-1]]
                rows = []
                for tl in table_lines[2:]:
                    rows.append([c.strip() for c in tl.split("|")[1:-1]])

                num_cols = len(headers)
                table = doc.add_table(rows=1 + len(rows), cols=num_cols)
                table.style = "Light Grid Accent 1"

                for ci, h in enumerate(headers):
                    table.rows[0].cells[ci].text = h
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row[:num_cols]):
                        table.rows[ri + 1].cells[ci].text = cell

        # Bold text (**text**)
        elif line.startswith("**") and "**" in line[2:]:
            p = doc.add_paragraph()
            run = p.add_run(line.strip().strip("*"))
            run.bold = True

        # Regular paragraphs
        elif line.strip():
            # Clean markdown formatting
            clean = line.strip()
            # Remove bold markers for plain text
            clean = clean.replace("**", "")
            if clean:
                doc.add_paragraph(clean)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- Route Handlers ---------------------------------------------------------

def handle_whoami(event):
    """GET /api/grc-tools/whoami — return user identity from JWT."""
    cookies = parse_cookies(event.get("headers", {}))
    token = cookies.get("grc_session")
    if not token:
        return error_response("Not authenticated", 401)

    info = extract_user_info(token)
    if not info:
        return error_response("Invalid token", 401)

    return json_response({
        "user_id": info["user_id"],
        "email": info["email"],
        "name": info["name"],
        "idp": parse_idp(info["idp"]),
    })


def handle_get_session(event):
    """GET /api/grc-tools/session — return saved wizard state."""
    user_id = get_user_id(event)
    if not user_id:
        return error_response("Not authenticated", 401)

    state = get_session(user_id)
    if state:
        # Clean DynamoDB keys
        clean = {k: state[k] for k in state
                 if k not in ("user_id",)
                 and not k.startswith("_")}
        return json_response(clean)
    return json_response(None)


def handle_put_session(event):
    """PUT /api/grc-tools/session — save wizard state."""
    user_id = get_user_id(event)
    if not user_id:
        return error_response("Not authenticated", 401)

    try:
        body = json.loads(event.get("body", "{}"))
    except json.JSONDecodeError:
        return error_response("Invalid JSON", 400)

    # Merge with existing state
    existing = get_session(user_id) or {}
    existing.update(body)
    # Don't store user_id or system keys
    existing.pop("user_id", None)

    updated_at = put_session(user_id, existing)
    return json_response({"saved": True, "updated_at": updated_at})


def handle_templates(event):
    """GET /api/grc-tools/templates — return template manifest."""
    # Public-ish — still behind auth gate but no user-specific data
    return json_response(MANIFEST)


def handle_generate(event):
    """POST /api/grc-tools/generate — render policies and save to S3."""
    user_id = get_user_id(event)
    if not user_id:
        return error_response("Not authenticated", 401)

    try:
        body = json.loads(event.get("body", "{}"))
    except json.JSONDecodeError:
        return error_response("Invalid JSON", 400)

    policies = body.get("policies", [])
    variables = body.get("variables", {})
    org_name = body.get("org_name", variables.get("company_name", "Organization"))

    if not policies:
        return error_response("No policies specified", 400)

    timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H%M%SZ")
    generated = []

    for policy_id in policies:
        # Find template file
        template_info = None
        for t in MANIFEST.get("templates", []):
            if t["id"] == policy_id:
                template_info = t
                break

        if not template_info:
            generated.append({
                "policy_id": policy_id,
                "error": "Template not found",
            })
            continue

        try:
            tmpl = jinja_env.get_template(template_info["file"])

            # Add standard variables
            ctx = dict(variables)
            ctx.setdefault("company_name", org_name)
            ctx.setdefault("effective_date", datetime.now(timezone.utc).strftime("%B %d, %Y"))

            md_content = tmpl.render(**ctx)
        except TemplateNotFound:
            generated.append({
                "policy_id": policy_id,
                "error": "Template file not found",
            })
            continue
        except Exception as e:
            generated.append({
                "policy_id": policy_id,
                "error": f"Render error: {str(e)}",
            })
            continue

        # Generate PDF
        try:
            pdf_bytes = render_pdf(md_content)
        except Exception as e:
            pdf_bytes = None

        # Generate DOCX
        try:
            docx_bytes = render_docx(md_content)
        except Exception as e:
            docx_bytes = None

        # Upload to S3
        base_key = f"{S3_PREFIX}/{user_id}"
        pdf_key = f"{base_key}/pdf/{policy_id}-{timestamp}.pdf"
        docx_key = f"{base_key}/docx/{policy_id}-{timestamp}.docx"

        if pdf_bytes:
            s3.put_object(
                Bucket=S3_BUCKET,
                Key=pdf_key,
                Body=pdf_bytes,
                ContentType="application/pdf",
            )

        if docx_bytes:
            s3.put_object(
                Bucket=S3_BUCKET,
                Key=docx_key,
                Body=docx_bytes,
                ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        # Record in DynamoDB
        policy_record = {
            "user_id": user_id,
            "policy_id": policy_id,
            "generated_at": timestamp,
            "template_version": template_info.get("version", "1.0"),
            "variables": variables,
        }
        if pdf_bytes:
            policy_record["pdf_s3_key"] = pdf_key
        if docx_bytes:
            policy_record["docx_s3_key"] = docx_key

        dynamodb.put_item(
            TableName=POLICIES_TABLE,
            Item=serialize_policy_record(policy_record),
        )

        result = {
            "policy_id": policy_id,
            "status": "generated",
        }
        if pdf_bytes:
            result["pdf_url"] = f"/grc-tools/users/{user_id}/pdf/{policy_id}-{timestamp}.pdf"
        if docx_bytes:
            result["docx_url"] = f"/grc-tools/users/{user_id}/docx/{policy_id}-{timestamp}.docx"

        generated.append(result)

    return json_response({
        "generated": generated,
        "total": len(policies),
        "completed": sum(1 for g in generated if g.get("status") == "generated"),
    })


def handle_bundle(event):
    """POST /api/grc-tools/bundle — create ZIP of generated policies."""
    user_id = get_user_id(event)
    if not user_id:
        return error_response("Not authenticated", 401)

    try:
        body = json.loads(event.get("body", "{}"))
    except json.JSONDecodeError:
        return error_response("Invalid JSON", 400)

    fmt = body.get("format", "pdf")
    policy_ids = body.get("policies", [])

    if fmt not in ("pdf", "docx"):
        return error_response("Format must be 'pdf' or 'docx'", 400)

    if not policy_ids:
        # Get all policies for this user from DynamoDB
        resp = dynamodb.query(
            TableName=POLICIES_TABLE,
            KeyConditionExpression="user_id = :uid",
            ExpressionAttributeValues={":uid": {"S": user_id}},
        )
        policy_ids = [item["policy_id"]["S"] for item in resp.get("Items", [])]

    if not policy_ids:
        return error_response("No policies found to bundle", 404)

    timestamp = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H%M%SZ")
    buf = io.BytesIO()

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for pid in policy_ids:
            # Look up the latest generated version from DynamoDB
            resp = dynamodb.get_item(
                TableName=POLICIES_TABLE,
                Key={
                    "user_id": {"S": user_id},
                    "policy_id": {"S": pid},
                },
            )
            item = resp.get("Item", {})
            key_field = f"{fmt}_s3_key"
            s3_key = item.get(key_field, {}).get("S")

            if s3_key:
                try:
                    obj = s3.get_object(Bucket=S3_BUCKET, Key=s3_key)
                    filename = f"{pid}.{fmt}"
                    zf.writestr(filename, obj["Body"].read())
                except Exception:
                    pass  # skip missing files

    zip_bytes = buf.getvalue()
    zip_key = f"{S3_PREFIX}/{user_id}/bundles/all-policies-{timestamp}-{fmt}.zip"

    s3.put_object(
        Bucket=S3_BUCKET,
        Key=zip_key,
        Body=zip_bytes,
        ContentType="application/zip",
    )

    file_count = len(policy_ids)
    return json_response({
        "zip_url": f"/grc-tools/users/{user_id}/bundles/all-policies-{timestamp}-{fmt}.zip",
        "size_bytes": len(zip_bytes),
        "file_count": file_count,
    })


def handle_policies(event):
    """GET /api/grc-tools/policies — list user's generated policies."""
    user_id = get_user_id(event)
    if not user_id:
        return error_response("Not authenticated", 401)

    resp = dynamodb.query(
        TableName=POLICIES_TABLE,
        KeyConditionExpression="user_id = :uid",
        ExpressionAttributeValues={":uid": {"S": user_id}},
    )

    policies = []
    for item in resp.get("Items", []):
        pid = item.get("policy_id", {}).get("S", "")
        generated_at = item.get("generated_at", {}).get("S", "")
        pdf_key = item.get("pdf_s3_key", {}).get("S", "")
        docx_key = item.get("docx_s3_key", {}).get("S", "")

        entry = {
            "policy_id": pid,
            "generated_at": generated_at,
        }
        if pdf_key:
            entry["pdf_url"] = f"/grc-tools/users/{user_id}/pdf/{pdf_key.split('/')[-1]}"
        if docx_key:
            entry["docx_url"] = f"/grc-tools/users/{user_id}/docx/{docx_key.split('/')[-1]}"
        policies.append(entry)

    return json_response({"policies": policies})


def serialize_policy_record(record):
    """Convert policy record dict to DynamoDB item."""
    item = {}
    for k, v in record.items():
        if isinstance(v, str):
            item[k] = {"S": v}
        elif isinstance(v, dict):
            item[k] = {"S": json.dumps(v)}
    return item


# --- Route table ------------------------------------------------------------

ROUTES = {
    ("GET", "/api/grc-tools/whoami"): handle_whoami,
    ("GET", "/api/grc-tools/session"): handle_get_session,
    ("PUT", "/api/grc-tools/session"): handle_put_session,
    ("GET", "/api/grc-tools/templates"): handle_templates,
    ("POST", "/api/grc-tools/generate"): handle_generate,
    ("POST", "/api/grc-tools/bundle"): handle_bundle,
    ("GET", "/api/grc-tools/policies"): handle_policies,
}


# --- Main handler -----------------------------------------------------------

def handler(event, context):
    """API Gateway v2 proxy integration handler."""
    # Validate origin secret
    if not check_origin_secret(event):
        return error_response("Forbidden", 403)

    # Route
    route_key = event.get("routeKey", "")
    http_method = event.get("requestContext", {}).get("http", {}).get("method", "")
    path = event.get("requestContext", {}).get("http", {}).get("path", "")

    # Fallback: parse routeKey
    if not http_method and route_key:
        parts = route_key.split(" ", 1)
        if len(parts) == 2:
            http_method, path = parts

    route_key_tuple = (http_method, path)
    handler_fn = ROUTES.get(route_key_tuple)

    if not handler_fn:
        return error_response(f"No route for {http_method} {path}", 404)

    try:
        return handler_fn(event)
    except Exception as e:
        traceback.print_exc()
        return error_response(f"Internal error: {str(e)}", 500)
